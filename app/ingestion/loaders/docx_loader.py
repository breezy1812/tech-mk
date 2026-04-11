from pathlib import Path

from app.domain.schemas.rag import SourceDocument


class DocxLoader:
    supported_suffixes = {".docx"}

    def load(self, path: Path, docs_root: Path) -> SourceDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to load DOCX documents") from exc

        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return SourceDocument(
            file_name=path.name,
            relative_path=path.relative_to(docs_root).as_posix(),
            source_type="docx",
            content="\n\n".join(paragraphs),
        )
